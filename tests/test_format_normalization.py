from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from rule_engine.format_normalization import run_format_normalization


def _set_run_font(run, name: str) -> None:
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)


class FormatNormalizationTests(unittest.TestCase):
    def setUp(self) -> None:
        self.tmpdir = tempfile.TemporaryDirectory()
        self.addCleanup(self.tmpdir.cleanup)
        self.doc_path = Path(self.tmpdir.name) / "doc.docx"

    def _run(self, **overrides) -> list:
        config = {"enabled": True}
        config.update(overrides)
        return run_format_normalization(self.doc_path, document_name="doc.docx", config=config)

    def test_normalizes_fonts_to_dominant_but_keeps_symbols(self) -> None:
        doc = Document()
        paragraph = doc.add_paragraph()
        _set_run_font(paragraph.add_run("Texto principal del documento en Arial."), "Arial")
        _set_run_font(paragraph.add_run("Otro Arial."), "Arial")
        _set_run_font(paragraph.add_run("TNR"), "Times New Roman")
        _set_run_font(paragraph.add_run(""), "Symbol")
        doc.save(self.doc_path)

        self._run()
        normalized = Document(str(self.doc_path))
        fonts = []
        for run in normalized.paragraphs[0].runs:
            rpr = run._element.find(qn("w:rPr"))
            rfonts = rpr.find(qn("w:rFonts")) if rpr is not None else None
            fonts.append(rfonts.get(qn("w:ascii")) if rfonts is not None else None)
        self.assertEqual(fonts[2], "Arial")
        self.assertEqual(fonts[3], "Symbol")

    def test_flattens_nested_tables_preserving_text(self) -> None:
        doc = Document()
        outer = doc.add_table(rows=1, cols=1)
        cell = outer.cell(0, 0)
        cell.text = "Contenido externo"
        inner = cell.add_table(rows=1, cols=2)
        inner.cell(0, 0).text = "Recuerde observar"
        inner.cell(0, 1).text = "los riesgos"
        doc.save(self.doc_path)

        records = self._run()
        actions = {record.action: record.count for record in records}
        self.assertEqual(actions.get("nested_tables_flattened", 0), 1)
        flattened = Document(str(self.doc_path))
        self.assertEqual(len(flattened.tables), 1)
        cell_text = flattened.tables[0].cell(0, 0).text
        self.assertIn("Contenido externo", cell_text)
        self.assertIn("Recuerde observar", cell_text)
        self.assertIn("los riesgos", cell_text)
        self.assertFalse(flattened.tables[0].cell(0, 0).tables)

    def test_dry_run_does_not_write(self) -> None:
        doc = Document()
        paragraph = doc.add_paragraph()
        for fragment in ["a", "b", "c"]:
            _set_run_font(paragraph.add_run(fragment), "Arial")
        doc.save(self.doc_path)
        before = self.doc_path.read_bytes()

        records = run_format_normalization(
            self.doc_path, document_name="doc.docx", config={"enabled": True}, dry_run=True
        )
        self.assertTrue(any(record.count for record in records))
        self.assertFalse(any(record.applied for record in records))
        self.assertEqual(before, self.doc_path.read_bytes())


if __name__ == "__main__":
    unittest.main()
