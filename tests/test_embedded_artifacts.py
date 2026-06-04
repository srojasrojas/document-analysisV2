from __future__ import annotations

import tempfile
import unittest
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement

from rule_engine import embedded_artifacts
from rule_engine.embedded_artifacts import run_embedded_artifact_cleanup


def _footer_xml_parts(path: Path) -> list[str]:
    with zipfile.ZipFile(path) as package:
        footer_names = sorted(name for name in package.namelist() if name.startswith("word/footer"))
        return [package.read(name).decode("utf-8", errors="ignore") for name in footer_names]


def _document_xml(path: Path) -> str:
    with zipfile.ZipFile(path) as package:
        return package.read("word/document.xml").decode("utf-8", errors="ignore")


def _footer_table_text(doc: Document) -> str:
    parts: list[str] = []
    for table in doc.sections[0].footer.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text.strip())
    return "\n".join(parts)


def _story_table_texts(tables: list) -> list[str]:
    values: list[str] = []
    for table in tables:
        parts: list[str] = []
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text.strip())
        values.append("\n".join(parts))
    return values


def _header_table_text(doc: Document) -> str:
    return "\n".join(_story_table_texts(list(doc.sections[0].header.tables)))


def _body_table_texts(doc: Document) -> list[str]:
    return _story_table_texts(list(doc.tables))


def _add_embedded_header_table(doc: Document, code: str = "P-PRMI-OM-101") -> None:
    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "Gerencia Producción Mina"
    table.cell(0, 1).text = code
    table.cell(1, 0).text = "Operación Camión de Extracción"
    table.cell(1, 1).text = "Operación Camión de Extracción"
    table.cell(2, 0).text = "Elaboró: Constanza Delgado Supervisor Mina"
    table.cell(2, 1).text = "Revisó: Juan Pablo Ramos Coordinador Mina Aprobó: Rodrigo Puelles"


def _make_tc(text: str):
    tc = OxmlElement("w:tc")
    paragraph = OxmlElement("w:p")
    run = OxmlElement("w:r")
    value = OxmlElement("w:t")
    value.text = text
    run.append(value)
    paragraph.append(run)
    tc.append(paragraph)
    return tc


class _BrokenRow:
    def __init__(self, *texts: str) -> None:
        self._tr = OxmlElement("w:tr")
        for text in texts:
            self._tr.append(_make_tc(text))

    @property
    def cells(self):
        raise ValueError("no `tc` element at grid_offset=8")


class _BrokenTable:
    def __init__(self, *rows: _BrokenRow) -> None:
        self.rows = list(rows)


class EmbeddedArtifactCleanupTests(unittest.TestCase):
    def test_table_text_falls_back_for_malformed_rows(self) -> None:
        table = _BrokenTable(_BrokenRow("Código", "P-PRMI-OM-146"), _BrokenRow("Versión", "29"))

        self.assertEqual("Código | P-PRMI-OM-146 | Versión | 29", embedded_artifacts._table_text(table))
        self.assertEqual(
            ["codigo", "p-prmi-om-146", "version", "29"],
            embedded_artifacts._table_cell_norms(table),
        )

    def test_removes_repeated_footer_paragraphs_and_writes_real_footer(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / "sample.docx"
            doc = Document()
            doc.add_paragraph("El operador planta debe informar al supervisor de turno cualquier desviacion critica.")
            for page in range(1, 4):
                doc.add_paragraph("Versión")
                doc.add_paragraph("1")
                doc.add_paragraph('"Este es un documento controlado"')
                doc.add_paragraph("Fecha de Autorización Junio 2023")
                doc.add_paragraph(f"Página {page} de 3")
                doc.add_paragraph("Próxima Revisión Junio 2026")
            doc.save(path)

            records = run_embedded_artifact_cleanup(
                path,
                document_name="sample.docx",
                config={"action": "remove", "enabled": True},
                dry_run=False,
            )

            cleaned = Document(str(path))
            full_text = "\n".join(paragraph.text for paragraph in cleaned.paragraphs)
            self.assertIn("operador planta", full_text)
            self.assertNotIn("documento controlado", full_text)
            self.assertNotIn("Página 2 de 3", full_text)
            footer_text = _footer_table_text(cleaned)
            footer_xml = "\n".join(_footer_xml_parts(path))
            self.assertIn("Versión", footer_text)
            self.assertIn("documento controlado", footer_text)
            self.assertRegex(footer_xml, r">\s*PAGE\s*<")
            self.assertRegex(footer_xml, r">\s*NUMPAGES\s*<")
            self.assertNotIn("Este es un documento controlado", _document_xml(path))
            self.assertEqual(1, sum(1 for record in records if record.action == "write_footer" and record.applied))
            self.assertGreaterEqual(sum(1 for record in records if record.applied), 15)

    def test_second_run_is_idempotent(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / "sample.docx"
            doc = Document()
            for page in range(1, 3):
                doc.add_paragraph("Versión")
                doc.add_paragraph("1")
                doc.add_paragraph('"Este es un documento controlado"')
                doc.add_paragraph(f"Página {page} de 2")
                doc.add_paragraph("Próxima Revisión Junio 2026")
            doc.save(path)

            run_embedded_artifact_cleanup(path, config={"action": "remove", "enabled": True}, dry_run=False)
            second = run_embedded_artifact_cleanup(path, config={"action": "remove", "enabled": True}, dry_run=False)

            self.assertEqual(0, sum(1 for record in second if record.action in {"remove", "clear_text"}))
            self.assertEqual(0, sum(1 for record in second if record.action == "write_footer"))

    def test_clean_document_does_not_receive_footer(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / "sample.docx"
            doc = Document()
            doc.add_paragraph("El operador planta debe informar al supervisor de turno cualquier desviacion critica.")
            doc.save(path)

            records = run_embedded_artifact_cleanup(path, config={"action": "remove", "enabled": True}, dry_run=False)

            self.assertEqual([], records)
            self.assertEqual([], _footer_xml_parts(path))

    def test_removes_split_page_number_fragments(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / "sample.docx"
            doc = Document()
            for page in range(1, 4):
                doc.add_paragraph("Versión 21")
                doc.add_paragraph('"Este es un documento controlado"')
                doc.add_paragraph("Fecha de Autorización Enero 2025")
                doc.add_paragraph("Página")
                doc.add_paragraph(f"{page} de 105")
                doc.add_paragraph("Próxima Revisión Enero 2027")
            doc.save(path)

            records = run_embedded_artifact_cleanup(path, config={"action": "remove", "enabled": True}, dry_run=False)
            cleaned = Document(str(path))
            full_text = "\n".join(paragraph.text for paragraph in cleaned.paragraphs)
            footer_text = _footer_table_text(cleaned)

            self.assertNotIn("Página", full_text)
            self.assertNotIn("2 de 105", full_text)
            self.assertIn("Página 1 de 105", footer_text)
            self.assertGreaterEqual(sum(1 for record in records if record.action == "remove"), 15)

    def test_existing_footer_is_protected_by_default(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / "sample.docx"
            doc = Document()
            doc.sections[0].footer.paragraphs[0].text = "Pie de pagina real existente"
            for page in range(1, 3):
                doc.add_paragraph("Versión")
                doc.add_paragraph("21")
                doc.add_paragraph('"Este es un documento controlado"')
                doc.add_paragraph("Fecha de Autorización Enero 2025")
                doc.add_paragraph(f"Página {page} de 2")
                doc.add_paragraph("Próxima Revisión Enero 2027")
            doc.save(path)

            records = run_embedded_artifact_cleanup(path, config={"action": "remove", "enabled": True}, dry_run=False)
            cleaned = Document(str(path))

            self.assertEqual(1, sum(1 for record in records if record.action == "footer_protected_existing"))
            self.assertIn("Pie de pagina real existente", cleaned.sections[0].footer.paragraphs[0].text)
            self.assertEqual(0, sum(1 for record in records if record.action == "write_footer"))

    def test_existing_footer_can_be_overwritten_when_enabled(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / "sample.docx"
            doc = Document()
            doc.sections[0].footer.paragraphs[0].text = "Pie de pagina real existente"
            for page in range(1, 3):
                doc.add_paragraph("Versión 18")
                doc.add_paragraph('"Este es un documento controlado"')
                doc.add_paragraph("Fecha de Autorización Noviembre 2024")
                doc.add_paragraph(f"Página {page} de 2")
                doc.add_paragraph("Próxima Revisión Noviembre 2026")
            doc.save(path)

            records = run_embedded_artifact_cleanup(
                path,
                config={"action": "remove", "enabled": True, "overwrite_existing_footer": True},
                dry_run=False,
            )
            cleaned = Document(str(path))
            footer_text = _footer_table_text(cleaned)
            footer_xml = "\n".join(_footer_xml_parts(path))

            self.assertEqual(1, sum(1 for record in records if record.action == "write_footer" and record.applied))
            self.assertNotIn("Pie de pagina real existente", footer_text)
            self.assertIn("Versión", footer_text)
            self.assertRegex(footer_xml, r">\s*PAGE\s*<")
            self.assertRegex(footer_xml, r">\s*NUMPAGES\s*<")

    def test_operational_text_is_not_flagged(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / "sample.docx"
            doc = Document()
            doc.add_paragraph("Página de procedimiento: el operador debe revisar el estado del equipo antes de iniciar.")
            doc.add_paragraph("Supervisor de Producción debe autorizar la continuidad operacional.")
            doc.save(path)

            records = run_embedded_artifact_cleanup(path, config={"action": "preview", "enabled": True}, dry_run=True)

            self.assertEqual([], records)

    def test_repeated_metadata_tables_are_reviewed_by_default(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / "sample.docx"
            doc = Document()
            for _ in range(2):
                table = doc.add_table(rows=3, cols=3)
                table.cell(0, 0).text = "Código P-PRPL-OC-506"
                table.cell(0, 1).text = "VERSIÓN"
                table.cell(0, 2).text = "Página 1 de 2"
                table.cell(1, 0).text = "Elaboró:"
                table.cell(1, 1).text = "Revisó:"
                table.cell(1, 2).text = "Aprobó:"
                table.cell(2, 0).text = "Fecha de Autorización Junio 2023"
            doc.save(path)

            records = run_embedded_artifact_cleanup(path, config={"action": "remove", "enabled": True}, dry_run=False)
            table_records = [record for record in records if record.block_type == "table"]

            self.assertEqual(2, len(table_records))
            self.assertIn("protected", {record.action for record in table_records})
            self.assertIn("review", {record.action for record in table_records})

    def test_repeated_embedded_header_tables_are_moved_to_real_header(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / "sample.docx"
            doc = Document()
            for _ in range(3):
                _add_embedded_header_table(doc)
            doc.add_paragraph("El operador planta debe informar al supervisor de turno cualquier desviacion critica.")
            doc.save(path)

            records = run_embedded_artifact_cleanup(
                path,
                document_name="sample.docx",
                config={"action": "remove", "enabled": True, "remove_header_table_artifacts": True},
                dry_run=False,
            )

            cleaned = Document(str(path))
            header_text = _header_table_text(cleaned)
            body_header_tables = [text for text in _body_table_texts(cleaned) if "P-PRMI-OM-101" in text]

            self.assertIn("P-PRMI-OM-101", header_text)
            self.assertEqual([], body_header_tables)
            self.assertEqual(1, sum(1 for record in records if record.action == "write_header" and record.applied))
            self.assertEqual(1, sum(1 for record in records if record.action == "move_to_header" and record.applied))
            self.assertEqual(2, sum(1 for record in records if record.action == "remove_table" and record.applied))

    def test_existing_header_protects_embedded_header_source(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / "sample.docx"
            doc = Document()
            doc.sections[0].header.paragraphs[0].text = "Encabezado real existente"
            for _ in range(3):
                _add_embedded_header_table(doc)
            doc.save(path)

            records = run_embedded_artifact_cleanup(
                path,
                document_name="sample.docx",
                config={"action": "remove", "enabled": True, "remove_header_table_artifacts": True},
                dry_run=False,
            )
            cleaned = Document(str(path))

            self.assertIn("Encabezado real existente", cleaned.sections[0].header.paragraphs[0].text)
            self.assertEqual(1, sum(1 for record in records if record.action == "header_protected_existing"))
            self.assertEqual(0, sum(1 for record in records if record.action == "write_header"))
            self.assertEqual(0, sum(1 for record in records if record.action == "move_to_header"))
            self.assertEqual(1, sum(1 for text in _body_table_texts(cleaned) if "P-PRMI-OM-101" in text))

    def test_single_embedded_header_table_stays_protected(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / "sample.docx"
            doc = Document()
            _add_embedded_header_table(doc)
            doc.save(path)

            records = run_embedded_artifact_cleanup(
                path,
                document_name="sample.docx",
                config={"action": "remove", "enabled": True, "remove_header_table_artifacts": True},
                dry_run=False,
            )
            cleaned = Document(str(path))

            self.assertEqual(
                1,
                sum(
                    1
                    for record in records
                    if record.block_type == "header_table" and record.action == "protected"
                ),
            )
            self.assertEqual(0, sum(1 for record in records if record.action in {"write_header", "move_to_header"}))
            self.assertEqual(1, sum(1 for text in _body_table_texts(cleaned) if "P-PRMI-OM-101" in text))

    def test_header_cleanup_is_idempotent(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / "sample.docx"
            doc = Document()
            for _ in range(3):
                _add_embedded_header_table(doc)
            doc.save(path)

            config = {"action": "remove", "enabled": True, "remove_header_table_artifacts": True}
            run_embedded_artifact_cleanup(path, document_name="sample.docx", config=config, dry_run=False)
            second = run_embedded_artifact_cleanup(path, document_name="sample.docx", config=config, dry_run=False)

            self.assertEqual(0, sum(1 for record in second if record.action == "write_header"))
            self.assertEqual(0, sum(1 for record in second if record.action == "move_to_header"))
            self.assertEqual(0, sum(1 for record in second if record.action == "remove_table"))


if __name__ == "__main__":
    unittest.main()
