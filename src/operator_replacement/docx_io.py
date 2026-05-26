from __future__ import annotations

import difflib
import re
import unicodedata
from collections.abc import Iterable
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

from .models import DocxElement


HEADING_STYLES = {"heading 1", "heading 2", "heading 3", "heading 4", "title", "subtitle"}


def normalize_text(text: str) -> str:
    normalized = unicodedata.normalize("NFKD", text or "")
    normalized = "".join(ch for ch in normalized if not unicodedata.combining(ch))
    normalized = normalized.lower().replace("\u00a0", " ")
    normalized = re.sub(r"\s+", " ", normalized)
    return normalized.strip()


def load_docx(path: str | Path) -> DocxDocument:
    return Document(str(path))


def save_docx(doc: DocxDocument, path: str | Path) -> None:
    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


def is_heading(paragraph: Paragraph) -> bool:
    return (paragraph.style.name or "").lower() in HEADING_STYLES


def _iter_cell_paragraphs(cell: _Cell, table_location: str) -> Iterable[tuple[str, Paragraph]]:
    for paragraph_index, paragraph in enumerate(cell.paragraphs):
        yield f"{table_location}:p{paragraph_index}", paragraph

    for nested_index, nested_table in enumerate(cell.tables):
        nested_prefix = f"{table_location}:table[{nested_index}]"
        for row_index, row in enumerate(nested_table.rows):
            for cell_index, nested_cell in enumerate(row.cells):
                nested_location = f"{nested_prefix}[{row_index}][{cell_index}]"
                yield from _iter_cell_paragraphs(nested_cell, nested_location)


def collect_elements(doc: DocxDocument) -> list[DocxElement]:
    elements: list[DocxElement] = []

    for paragraph_index, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text
        if not text.strip():
            continue
        elements.append(
            DocxElement(
                text=text,
                normalized=normalize_text(text),
                location=f"body:{paragraph_index}",
                paragraph_obj=paragraph,
                is_heading=is_heading(paragraph),
            )
        )

    for table_index, table in enumerate(doc.tables):
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                cell_location = f"table[{table_index}][{row_index}][{cell_index}]"
                for location, paragraph in _iter_cell_paragraphs(cell, cell_location):
                    text = paragraph.text
                    if not text.strip():
                        continue
                    elements.append(
                        DocxElement(
                            text=text,
                            normalized=normalize_text(text),
                            location=location,
                            paragraph_obj=paragraph,
                            is_heading=is_heading(paragraph),
                        )
                    )

    return elements


def _fallback_replace(paragraph: Paragraph, new_text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(new_text)


def apply_surgical_change(element: DocxElement, new_text: str) -> None:
    paragraph = element.paragraph_obj
    if not paragraph.runs:
        paragraph.add_run(new_text)
        element.text = new_text
        element.normalized = normalize_text(new_text)
        return

    original_full = "".join(run.text for run in paragraph.runs)
    if original_full == new_text:
        return

    run_spans: list[tuple[int, int]] = []
    cursor = 0
    for run in paragraph.runs:
        run_length = len(run.text)
        run_spans.append((cursor, cursor + run_length))
        cursor += run_length

    matcher = difflib.SequenceMatcher(None, original_full, new_text, autojunk=False)
    changed_original_start = len(original_full)
    changed_original_end = 0
    changed_new_start = len(new_text)
    changed_new_end = 0
    has_change = False

    for tag, original_start, original_end, new_start, new_end in matcher.get_opcodes():
        if tag == "equal":
            continue
        changed_original_start = min(changed_original_start, original_start)
        changed_original_end = max(changed_original_end, original_end)
        changed_new_start = min(changed_new_start, new_start)
        changed_new_end = max(changed_new_end, new_end)
        has_change = True

    if not has_change:
        return

    affected = [
        index
        for index, (run_start, run_end) in enumerate(run_spans)
        if run_start < changed_original_end and run_end > changed_original_start
    ]
    if not affected:
        _fallback_replace(paragraph, new_text)
        element.text = new_text
        element.normalized = normalize_text(new_text)
        return

    first_index = affected[0]
    last_index = affected[-1]
    first_run_start = run_spans[first_index][0]
    last_run_end = run_spans[last_index][1]
    prefix_length = max(0, changed_original_start - first_run_start)
    suffix_length = max(0, last_run_end - changed_original_end)
    replacement = new_text[changed_new_start:changed_new_end]

    first_run = paragraph.runs[first_index]
    last_run = paragraph.runs[last_index]
    prefix = first_run.text[:prefix_length]
    suffix = last_run.text[len(last_run.text) - suffix_length :] if suffix_length else ""

    first_run.text = prefix + replacement + suffix
    for index in affected[1:]:
        paragraph.runs[index].text = ""

    element.text = new_text
    element.normalized = normalize_text(new_text)