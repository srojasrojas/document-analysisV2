from __future__ import annotations

import difflib
import re
import unicodedata
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

from .models import DocxElement


HEADING_STYLES = {"heading 1", "heading 2", "heading 3", "heading 4", "title", "subtitle"}
ACTION_SECTION_HINTS = {
    "alcance",
    "definiciones",
    "desarrollo",
    "descripcion",
    "ejecucion",
    "instrucciones",
    "objetivo",
    "procedimiento",
    "registros",
    "responsabilidades",
    "roles",
}


def _is_registros_title(text: str) -> bool:
    normalized = normalize_text(text)
    normalized = re.sub(r"^\d+(?:\.\d+)*\s*[-.)]?\s*", "", normalized)
    return bool(re.fullmatch(r"(?:anexos?\s+)?registros?", normalized))


def normalize_text(text: str) -> str:
    normalized = unicodedata.normalize("NFKD", text or "")
    normalized = "".join(ch for ch in normalized if not unicodedata.combining(ch))
    normalized = normalized.lower().replace("\u00a0", " ")
    normalized = re.sub(r"\s+", " ", normalized)
    return normalized.strip()


def load_docx(path: str | Path) -> DocxDocument:
    return Document(str(path))


def save_docx(doc: DocxDocument, path: str | Path) -> None:
    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


def is_heading(paragraph: Paragraph) -> bool:
    return (paragraph.style.name or "").lower() in HEADING_STYLES


def _looks_like_section_title(text: str, paragraph: Paragraph | None = None) -> bool:
    stripped = re.sub(r"\s+", " ", text or "").strip(" .:\t\r\n")
    if not stripped or len(stripped) > 120:
        return False
    if paragraph is not None and is_heading(paragraph):
        return True

    normalized = normalize_text(stripped)
    without_number = re.sub(r"^\d+(?:\.\d+)*\s*[-.)]?\s*", "", normalized)
    if without_number in ACTION_SECTION_HINTS:
        return True
    if _is_registros_title(without_number):
        return True

    letters = [char for char in stripped if char.isalpha()]
    if letters and sum(char.isupper() for char in letters) / len(letters) >= 0.8:
        return len(stripped.split()) <= 10
    return False


def _is_excluded_section(section_path: tuple[str, ...]) -> bool:
    return any(_is_registros_title(section) for section in section_path)


def _append_element(
    elements: list[DocxElement],
    *,
    paragraph: Paragraph,
    location: str,
    block_type: str,
    section_path: tuple[str, ...],
    table_index: int | None = None,
    row_index: int | None = None,
    cell_index: int | None = None,
) -> None:
    text = paragraph.text
    if not text.strip():
        return
    paragraph_is_heading = is_heading(paragraph) or _looks_like_section_title(text, paragraph)
    element_section_path = section_path
    if paragraph_is_heading and not section_path or (
        paragraph_is_heading and normalize_text(text) not in {normalize_text(section) for section in section_path}
    ):
        element_section_path = (*section_path, re.sub(r"\s+", " ", text).strip())

    elements.append(
        DocxElement(
            text=text,
            normalized=normalize_text(text),
            location=location,
            paragraph_obj=paragraph,
            is_heading=paragraph_is_heading,
            block_type=block_type,
            section_path=element_section_path,
            in_excluded_section=_is_excluded_section(element_section_path),
            table_index=table_index,
            row_index=row_index,
            cell_index=cell_index,
        )
    )


def _iter_table_elements(
    *,
    doc: DocxDocument,
    table_element,
    table_location: str,
    table_index: int | None,
    section_path: tuple[str, ...],
) -> list[DocxElement]:
    elements: list[DocxElement] = []
    local_section_path = section_path

    for row_index, row in enumerate(table_element.findall(qn("w:tr"))):
        row_section_path = local_section_path
        cells = row.findall(qn("w:tc"))
        for cell_index, cell in enumerate(cells):
            cell_location = f"{table_location}[{row_index}][{cell_index}]"
            paragraph_index = 0
            nested_table_index = 0
            for child in cell.iterchildren():
                if child.tag == qn("w:p"):
                    paragraph = Paragraph(child, doc)
                    text = paragraph.text
                    if _looks_like_section_title(text, paragraph):
                        row_section_path = (*local_section_path, re.sub(r"\s+", " ", text).strip())
                    _append_element(
                        elements,
                        paragraph=paragraph,
                        location=f"{cell_location}:p{paragraph_index}",
                        block_type="table",
                        section_path=row_section_path,
                        table_index=table_index,
                        row_index=row_index,
                        cell_index=cell_index,
                    )
                    paragraph_index += 1
                elif child.tag == qn("w:tbl"):
                    nested_location = f"{cell_location}:table[{nested_table_index}]"
                    elements.extend(
                        _iter_table_elements(
                            doc=doc,
                            table_element=child,
                            table_location=nested_location,
                            table_index=table_index,
                            section_path=row_section_path,
                        )
                    )
                    nested_table_index += 1
        if _is_excluded_section(row_section_path):
            local_section_path = row_section_path
    return elements


def collect_elements(doc: DocxDocument) -> list[DocxElement]:
    elements: list[DocxElement] = []
    section_path: tuple[str, ...] = ()
    table_index = 0
    paragraph_index = 0

    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            paragraph = Paragraph(child, doc)
            if _looks_like_section_title(paragraph.text, paragraph):
                section_path = (re.sub(r"\s+", " ", paragraph.text).strip(),)
            _append_element(
                elements,
                paragraph=paragraph,
                location=f"body:{paragraph_index}",
                block_type="body",
                section_path=section_path,
            )
            paragraph_index += 1
        elif child.tag == qn("w:tbl"):
            elements.extend(
                _iter_table_elements(
                    doc=doc,
                    table_element=child,
                    table_location=f"table[{table_index}]",
                    table_index=table_index,
                    section_path=section_path,
                )
            )
            table_index += 1

    return elements


def _fallback_replace(paragraph: Paragraph, new_text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(new_text)


def apply_surgical_change(element: DocxElement, new_text: str) -> None:
    paragraph = element.paragraph_obj
    if not paragraph.runs:
        paragraph.add_run(new_text)
        element.text = new_text
        element.normalized = normalize_text(new_text)
        return

    original_full = "".join(run.text for run in paragraph.runs)
    if original_full == new_text:
        return

    run_spans: list[tuple[int, int]] = []
    cursor = 0
    for run in paragraph.runs:
        run_length = len(run.text)
        run_spans.append((cursor, cursor + run_length))
        cursor += run_length

    matcher = difflib.SequenceMatcher(None, original_full, new_text, autojunk=False)
    changed_original_start = len(original_full)
    changed_original_end = 0
    changed_new_start = len(new_text)
    changed_new_end = 0
    has_change = False

    for tag, original_start, original_end, new_start, new_end in matcher.get_opcodes():
        if tag == "equal":
            continue
        changed_original_start = min(changed_original_start, original_start)
        changed_original_end = max(changed_original_end, original_end)
        changed_new_start = min(changed_new_start, new_start)
        changed_new_end = max(changed_new_end, new_end)
        has_change = True

    if not has_change:
        return

    affected = [
        index
        for index, (run_start, run_end) in enumerate(run_spans)
        if run_start < changed_original_end and run_end > changed_original_start
    ]
    if not affected:
        _fallback_replace(paragraph, new_text)
        element.text = new_text
        element.normalized = normalize_text(new_text)
        return

    first_index = affected[0]
    last_index = affected[-1]
    first_run_start = run_spans[first_index][0]
    last_run_end = run_spans[last_index][1]
    prefix_length = max(0, changed_original_start - first_run_start)
    suffix_length = max(0, last_run_end - changed_original_end)
    replacement = new_text[changed_new_start:changed_new_end]

    first_run = paragraph.runs[first_index]
    last_run = paragraph.runs[last_index]
    prefix = first_run.text[:prefix_length]
    suffix = last_run.text[len(last_run.text) - suffix_length :] if suffix_length else ""

    first_run.text = prefix + replacement + suffix
    for index in affected[1:]:
        paragraph.runs[index].text = ""

    element.text = new_text
    element.normalized = normalize_text(new_text)