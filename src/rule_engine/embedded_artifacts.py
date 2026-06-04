from __future__ import annotations

import hashlib
import re
from collections import Counter
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from docx.document import Document as DocxDocument
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph

from .docx_io import collect_elements, load_docx, normalize_text, save_docx
from .models import EmbeddedArtifactRecord


ACTION_VERB_RE = re.compile(
    r"\b(?:debe(?:n)?|deber[aá]n?|realiza(?:r|n)?|revisa(?:r|n)?|verifica(?:r|n)?|"
    r"detiene(?:r|n)?|opera(?:r|n)?|inspecciona(?:r|n)?|coordina(?:r|n)?|informa(?:r|n)?|"
    r"avisa(?:r|n)?|solicita(?:r|n)?|autoriza(?:r|n)?|registra(?:r|n)?|bloquea(?:r|n)?|"
    r"desbloquea(?:r|n)?|energiza(?:r|n)?|desenergiza(?:r|n)?|comunica(?:r|n)?|asegura(?:r|n)?)\b",
    re.IGNORECASE,
)

MONTH_RE = r"(?:enero|febrero|marzo|abril|mayo|junio|julio|agosto|septiembre|setiembre|octubre|noviembre|diciembre)"
TEXT_CHILD_TAGS = {
    qn("w:t"),
    qn("w:delText"),
    qn("w:instrText"),
    qn("w:tab"),
    qn("w:noBreakHyphen"),
    qn("w:softHyphen"),
}

DEFAULT_CLEANUP_CONFIG: dict[str, Any] = {
    "enabled": False,
    "action": "preview",
    "apply_before_rules": True,
    "min_confidence_remove": 0.82,
    "min_confidence_review": 0.45,
    "max_removable_chars": 160,
    "remove_table_artifacts": False,
    "protect_front_matter_tables": True,
    "front_matter_table_count": 1,
    "write_real_footer": True,
    "overwrite_existing_footer": False,
    "footer_table_width_inches": 7.0,
    "footer_font_size_pt": 8,
    "footer_write_first_even_variants": True,
    "patterns": [],
    "table_patterns": [],
    "remove_header_table_artifacts": False,
    "header_table_front_matter_count": 1,
    "write_real_header": True,
    "overwrite_existing_header": False,
}

DEFAULT_PATTERNS = [
    r"\beste\s+es\s+un\s+documento\s+controlado\b",
    r"\bfecha\s+de\s+autorizacion\b",
    r"\bproxima\s+revision\b",
    r"\bpagina\s+\d+(?:\s+de\s+\d+)?\b",
]

DEFAULT_TABLE_PATTERNS = [
    r"\bcodigo\b",
    r"\bversion\b",
    r"\belaboro\b",
    r"\breviso\b",
    r"\baprobo\b",
    r"\bfecha\s+de\s+autorizacion\b",
]

_IMAGE_RELTYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image"
_BLIP_QNAME = "{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
_R_EMBED_QNAME = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
_HEADER_APPROVAL_RE = re.compile(r"\b(elaboro|reviso|aprobo)\b")
_HEADER_DOC_CODE_RE = re.compile(r"\bp-[a-z]+-[a-z]+-\d{3,}\b")


@dataclass
class _ParagraphInfo:
    paragraph: Paragraph
    index: int
    location: str
    text: str
    normalized: str
    style_name: str
    alignment: str
    section_path: tuple[str, ...]


@dataclass
class _TableInfo:
    table: Table
    index: int
    location: str
    text: str
    normalized: str
    signature: str


@dataclass
class _HeaderTableInfo:
    table: Table
    index: int
    location: str
    text: str
    normalized: str
    tbl_element: Any


@dataclass
class _ArtifactCandidate:
    record: EmbeddedArtifactRecord
    paragraph: Paragraph | None = None
    table: Table | None = None
    body_index: int | None = None
    table_index: int | None = None
    has_structural_break: bool = False


@dataclass
class _FooterMetadata:
    version: str = ""
    controlled_text: str = ""
    authorization_date: str = ""
    next_revision: str = ""
    page_total: str = ""
    source_locations: tuple[str, ...] = ()

    def is_sufficient(self) -> bool:
        strong_signals = sum(
            bool(value)
            for value in (
                self.controlled_text,
                self.authorization_date,
                self.next_revision,
                self.page_total,
            )
        )
        return strong_signals >= 2 and bool(self.page_total or self.controlled_text)

    def summary_text(self) -> str:
        parts: list[str] = []
        if self.version:
            parts.append(f"Versión {self.version}")
        if self.controlled_text:
            parts.append(self.controlled_text)
        if self.authorization_date:
            parts.append(self.authorization_date)
        parts.append("Página PAGE de NUMPAGES")
        if self.next_revision:
            parts.append(self.next_revision)
        return " | ".join(parts)


def resolve_cleanup_config(config: dict[str, Any] | None = None) -> dict[str, Any]:
    resolved = dict(DEFAULT_CLEANUP_CONFIG)
    for key, value in (config or {}).items():
        resolved[key] = value
    resolved["patterns"] = [*DEFAULT_PATTERNS, *list(resolved.get("patterns") or [])]
    resolved["table_patterns"] = [*DEFAULT_TABLE_PATTERNS, *list(resolved.get("table_patterns") or [])]
    return resolved


def run_embedded_artifact_cleanup(
    doc_path: str | Path,
    *,
    document_name: str | None = None,
    config: dict[str, Any] | None = None,
    dry_run: bool = True,
) -> list[EmbeddedArtifactRecord]:
    cleanup_config = resolve_cleanup_config(config)
    document_label = document_name or Path(doc_path).name
    doc = load_docx(doc_path)
    candidates = _scan_candidates(doc, document_label, cleanup_config)

    should_apply = cleanup_config.get("action") == "remove" and not dry_run
    header_record, header_changed = _header_reconstruction_record(
        doc,
        candidates,
        document_name=document_label,
        config=cleanup_config,
        apply_changes=should_apply,
    )
    footer_record, footer_changed = _footer_reconstruction_record(
        doc,
        candidates,
        document_name=document_label,
        config=cleanup_config,
        apply_changes=should_apply,
    )
    changed = False
    if should_apply:
        changed = _apply_candidates(candidates)
        if header_changed or footer_changed:
            changed = True
        if changed:
            save_docx(doc, doc_path)
    records = [candidate.record for candidate in candidates]
    if header_record is not None:
        records.append(header_record)
    if footer_record is not None:
        records.append(footer_record)
    return records


def scan_embedded_artifacts(
    doc_path: str | Path,
    *,
    document_name: str | None = None,
    config: dict[str, Any] | None = None,
) -> list[EmbeddedArtifactRecord]:
    return run_embedded_artifact_cleanup(
        doc_path,
        document_name=document_name,
        config={**(config or {}), "action": "preview"},
        dry_run=True,
    )


def excluded_locations(records: list[EmbeddedArtifactRecord]) -> set[str]:
    return {record.location for record in records if record.action == "exclude"}


def _scan_candidates(
    doc: DocxDocument,
    document_name: str,
    config: dict[str, Any],
) -> list[_ArtifactCandidate]:
    paragraphs = _paragraph_infos(doc)
    paragraph_counts = Counter(info.normalized for info in paragraphs if info.normalized)
    tables = _table_infos(doc)
    table_counts = Counter(info.signature for info in tables if info.signature)
    real_header_footer_texts = _real_header_footer_texts(doc)

    candidates: list[_ArtifactCandidate] = []
    for info in paragraphs:
        candidate = _paragraph_candidate(
            info,
            paragraphs=paragraphs,
            document_name=document_name,
            occurrence_count=paragraph_counts[info.normalized],
            real_header_footer_texts=real_header_footer_texts,
            config=config,
        )
        if candidate:
            candidates.append(candidate)

    for info in tables:
        candidate = _table_candidate(
            info,
            document_name=document_name,
            occurrence_count=table_counts[info.signature],
            real_header_footer_texts=real_header_footer_texts,
            config=config,
        )
        if candidate:
            candidates.append(candidate)

    header_table_infos = _header_table_infos(tables)
    header_sig_counts = Counter(info.normalized for info in header_table_infos if info.normalized)
    for position, info in enumerate(header_table_infos):
        candidate = _header_table_candidate(
            info,
            document_name=document_name,
            occurrence_count=header_sig_counts[info.normalized],
            config=config,
            header_table_position=position,
        )
        if candidate:
            candidates.append(candidate)

    return candidates


def _paragraph_infos(doc: DocxDocument) -> list[_ParagraphInfo]:
    sections_by_location = {element.location: element.section_path for element in collect_elements(doc)}
    infos: list[_ParagraphInfo] = []
    for index, paragraph in enumerate(doc.paragraphs):
        text = _compact(paragraph.text)
        if not text:
            continue
        location = f"body:{index}"
        infos.append(
            _ParagraphInfo(
                paragraph=paragraph,
                index=index,
                location=location,
                text=text,
                normalized=normalize_text(text),
                style_name=(paragraph.style.name if paragraph.style is not None else ""),
                alignment=_alignment_name(paragraph),
                section_path=sections_by_location.get(location, ()),
            )
        )
    return infos


def _table_infos(doc: DocxDocument) -> list[_TableInfo]:
    infos: list[_TableInfo] = []
    for index, table in enumerate(doc.tables):
        text = _table_text(table)
        normalized = normalize_text(text)
        infos.append(
            _TableInfo(
                table=table,
                index=index,
                location=f"table[{index}]",
                text=text,
                normalized=normalized,
                signature=_table_signature(table),
            )
        )
    return infos


def _paragraph_candidate(
    info: _ParagraphInfo,
    *,
    paragraphs: list[_ParagraphInfo],
    document_name: str,
    occurrence_count: int,
    real_header_footer_texts: set[str],
    config: dict[str, Any],
) -> _ArtifactCandidate | None:
    reasons: list[str] = []
    confidence = 0.0
    normalized = info.normalized
    neighbor_norms = _neighbor_norms(paragraphs, info.index, radius=3)
    metadata_neighbor_count = sum(1 for text in neighbor_norms if _is_metadata_text(text, config))

    confidence = max(confidence, _score_metadata_text(normalized, reasons, config))

    if re.fullmatch(r"\d+(?:\.\d+){0,2}", normalized) and any(
        neighbor == "version" or neighbor.startswith("version ") for neighbor in neighbor_norms
    ):
        confidence = max(confidence, 0.86)
        reasons.append("numeric_version_neighbor")

    if normalized == "pagina" and any(re.fullmatch(r"\d+\s+de\s+\d+", neighbor) for neighbor in neighbor_norms):
        confidence = max(confidence, 0.88)
        reasons.append("page_number_footer")

    if re.fullmatch(r"\d+\s+de\s+\d+", normalized) and any(neighbor == "pagina" for neighbor in neighbor_norms):
        confidence = max(confidence, 0.88)
        reasons.append("page_number_footer_fragment")

    if len(normalized) <= 90 and metadata_neighbor_count >= 2:
        confidence = max(confidence, 0.74)
        reasons.append("near_footer_metadata_cluster")

    if occurrence_count >= 3 and _is_metadata_text(normalized, config):
        confidence += 0.14
        reasons.append("recurring_metadata_text")

    real_match = normalized in real_header_footer_texts
    if real_match:
        confidence += 0.2
        reasons.append("matches_real_header_or_footer")

    if info.alignment in {"CENTER", "RIGHT"} and len(normalized) <= 90 and _is_metadata_text(normalized, config):
        confidence += 0.04
        reasons.append("short_aligned_metadata")

    if ACTION_VERB_RE.search(info.text) and not _is_metadata_text(normalized, config):
        confidence -= 0.35
        reasons.append("protected_action_text")

    confidence = max(0.0, min(1.0, confidence))
    if confidence < float(config.get("min_confidence_review", 0.45)):
        return None

    has_structural_break = _has_structural_break(info.paragraph)
    action = _action_for_paragraph(info, confidence, has_structural_break, config, reasons)
    record = EmbeddedArtifactRecord(
        document_name=document_name,
        location=info.location,
        block_type="body",
        action=action,
        confidence=round(confidence, 3),
        reasons=_unique(reasons),
        text=info.text,
        normalized_text=normalized,
        style_name=info.style_name,
        alignment=info.alignment,
        section_path=info.section_path,
        occurrence_count=occurrence_count,
        recurring_group_id=_hash_id(normalized) if occurrence_count > 1 else "",
        real_header_footer_match=real_match,
        context_before=_context_before(paragraphs, info.index),
        context_after=_context_after(paragraphs, info.index),
        candidate_id=_candidate_id(document_name, info.location, info.text),
    )
    return _ArtifactCandidate(
        record=record,
        paragraph=info.paragraph,
        body_index=info.index,
        has_structural_break=has_structural_break,
    )


def _table_candidate(
    info: _TableInfo,
    *,
    document_name: str,
    occurrence_count: int,
    real_header_footer_texts: set[str],
    config: dict[str, Any],
) -> _ArtifactCandidate | None:
    if not info.normalized:
        return None
    reasons: list[str] = []
    label_hits = _table_label_hits(info.normalized, config)
    confidence = 0.0
    if label_hits >= 3:
        confidence = max(confidence, 0.55)
        reasons.append("metadata_table_labels")
    if occurrence_count >= 2 and label_hits >= 2:
        confidence += 0.27
        reasons.append("recurring_metadata_table")
    real_match = any(part in real_header_footer_texts for part in _table_cell_norms(info.table))
    if real_match:
        confidence += 0.12
        reasons.append("contains_real_header_or_footer_text")

    confidence = max(0.0, min(1.0, confidence))
    if confidence < float(config.get("min_confidence_review", 0.45)):
        return None

    action = _action_for_table(info, confidence, config, reasons)
    record = EmbeddedArtifactRecord(
        document_name=document_name,
        location=info.location,
        block_type="table",
        action=action,
        confidence=round(confidence, 3),
        reasons=_unique(reasons),
        text=info.text,
        normalized_text=info.normalized,
        table_index=info.index,
        occurrence_count=occurrence_count,
        recurring_group_id=_hash_id(info.signature) if occurrence_count > 1 else "",
        real_header_footer_match=real_match,
        candidate_id=_candidate_id(document_name, info.location, info.text),
    )
    return _ArtifactCandidate(record=record, table=info.table, table_index=info.index)


def _action_for_paragraph(
    info: _ParagraphInfo,
    confidence: float,
    has_structural_break: bool,
    config: dict[str, Any],
    reasons: list[str],
) -> str:
    action = str(config.get("action", "preview")).lower()
    min_remove = float(config.get("min_confidence_remove", 0.82))
    max_chars = int(config.get("max_removable_chars", 160))
    if "protected_action_text" in reasons or len(info.text) > max_chars:
        return "protected"
    if confidence < min_remove:
        return "review"
    if action == "exclude":
        return "exclude"
    cleanup_action = "clear_text" if has_structural_break else "remove"
    if action == "remove":
        return cleanup_action
    return f"would_{cleanup_action}"


def _action_for_table(
    info: _TableInfo,
    confidence: float,
    config: dict[str, Any],
    reasons: list[str],
) -> str:
    action = str(config.get("action", "preview")).lower()
    min_remove = float(config.get("min_confidence_remove", 0.82))
    remove_tables = bool(config.get("remove_table_artifacts", False))
    protect_front_matter = bool(config.get("protect_front_matter_tables", True))
    front_matter_count = int(config.get("front_matter_table_count", 1))

    if protect_front_matter and info.index < front_matter_count:
        reasons.append("protected_front_matter_table")
        return "protected"
    if not remove_tables:
        reasons.append("table_removal_disabled")
        return "review"
    if confidence < min_remove:
        return "review"
    if action == "remove":
        return "remove_table"
    if action == "exclude":
        return "exclude"
    return "would_remove_table"


def _apply_candidates(candidates: list[_ArtifactCandidate]) -> bool:
    changed = False
    for candidate in candidates:
        if candidate.record.action != "clear_text" or candidate.paragraph is None:
            continue
        _clear_paragraph_text_preserving_breaks(candidate.paragraph)
        candidate.record.applied = True
        changed = True

    for candidate in sorted(
        (item for item in candidates if item.record.action == "remove" and item.paragraph is not None),
        key=lambda item: item.body_index if item.body_index is not None else -1,
        reverse=True,
    ):
        _remove_paragraph(candidate.paragraph)
        candidate.record.applied = True
        changed = True

    for candidate in sorted(
        (
            item
            for item in candidates
            if item.record.action in {"remove_table", "move_to_header"} and item.table is not None
        ),
        key=lambda item: item.table_index if item.table_index is not None else -1,
        reverse=True,
    ):
        _remove_table(candidate.table)
        candidate.record.applied = True
        changed = True
    return changed


def _footer_reconstruction_record(
    doc: DocxDocument,
    candidates: list[_ArtifactCandidate],
    *,
    document_name: str,
    config: dict[str, Any],
    apply_changes: bool,
) -> tuple[EmbeddedArtifactRecord | None, bool]:
    if not bool(config.get("write_real_footer", True)):
        return None, False

    metadata = _extract_footer_metadata(candidates)
    if not metadata.is_sufficient():
        if not _has_footer_source_candidates(candidates):
            return None, False
        return (
            _make_footer_record(
                document_name=document_name,
                action="footer_write_skipped_insufficient_metadata",
                metadata=metadata,
                reasons=["insufficient_footer_metadata"],
                applied=False,
                confidence=0.0,
            ),
            False,
        )

    if _existing_true_footer_present(doc) and not bool(config.get("overwrite_existing_footer", False)):
        return (
            _make_footer_record(
                document_name=document_name,
                action="footer_protected_existing",
                metadata=metadata,
                reasons=["existing_true_footer_detected", "overwrite_existing_footer_disabled"],
                applied=False,
            ),
            False,
        )

    action = "write_footer" if apply_changes else "would_write_footer"
    if apply_changes:
        _write_real_footer(doc, metadata, config)
    return (
        _make_footer_record(
            document_name=document_name,
            action=action,
            metadata=metadata,
            reasons=["reconstructed_real_footer", "dynamic_page_fields"],
            applied=apply_changes,
        ),
        apply_changes,
    )


def _extract_footer_metadata(candidates: list[_ArtifactCandidate]) -> _FooterMetadata:
    version_values: list[str] = []
    controlled_values: list[str] = []
    authorization_values: list[str] = []
    next_revision_values: list[str] = []
    page_total_values: list[str] = []
    source_locations: list[str] = []

    body_candidates = [
        candidate
        for candidate in candidates
        if candidate.record.block_type == "body" and candidate.record.action != "protected"
    ]
    candidates_by_body_index = {
        candidate.body_index: candidate
        for candidate in body_candidates
        if candidate.body_index is not None
    }

    for candidate in body_candidates:
        record = candidate.record
        normalized = record.normalized_text
        if not _is_footer_source_record(record):
            continue
        source_locations.append(record.location)

        version_match = re.fullmatch(r"version\s+(\d+(?:\.\d+){0,2})", normalized)
        if version_match:
            version_values.append(version_match.group(1))
        elif normalized == "version" and candidate.body_index is not None:
            next_candidate = candidates_by_body_index.get(candidate.body_index + 1)
            if next_candidate and re.fullmatch(r"\d+(?:\.\d+){0,2}", next_candidate.record.normalized_text):
                version_values.append(next_candidate.record.text)
        elif "numeric_version_neighbor" in record.reasons and re.fullmatch(r"\d+(?:\.\d+){0,2}", normalized):
            version_values.append(record.text)

        if "este es un documento controlado" in normalized:
            controlled_values.append(record.text)
        if "fecha de autorizacion" in normalized:
            authorization_values.append(record.text)
        if "proxima revision" in normalized:
            next_revision_values.append(record.text)
        page_match = re.fullmatch(r"pagina\s+\d+(?:\s+de\s+(\d+))?", normalized)
        if page_match and page_match.group(1):
            page_total_values.append(page_match.group(1))
        page_fragment_match = re.fullmatch(r"\d+\s+de\s+(\d+)", normalized)
        if page_fragment_match and "page_number_footer_fragment" in record.reasons:
            page_total_values.append(page_fragment_match.group(1))

    return _FooterMetadata(
        version=_most_common_text(version_values),
        controlled_text=_most_common_text(controlled_values),
        authorization_date=_most_common_text(authorization_values),
        next_revision=_most_common_text(next_revision_values),
        page_total=_largest_number_text(page_total_values),
        source_locations=tuple(_unique(source_locations)),
    )


def _has_footer_source_candidates(candidates: list[_ArtifactCandidate]) -> bool:
    return any(_is_footer_source_record(candidate.record) for candidate in candidates)


def _is_footer_source_record(record: EmbeddedArtifactRecord) -> bool:
    if record.block_type != "body" or record.action == "protected":
        return False
    footer_reasons = {
        "authorization_date_footer",
        "controlled_document_footer",
        "footer_label_version",
        "next_revision_footer",
        "numeric_version_neighbor",
        "page_number_footer",
        "page_number_footer_fragment",
    }
    return bool(footer_reasons.intersection(record.reasons))


def _make_footer_record(
    *,
    document_name: str,
    action: str,
    metadata: _FooterMetadata,
    reasons: list[str],
    applied: bool,
    confidence: float = 1.0,
) -> EmbeddedArtifactRecord:
    text = metadata.summary_text()
    return EmbeddedArtifactRecord(
        document_name=document_name,
        location="footer:default",
        block_type="footer",
        action=action,
        confidence=confidence,
        reasons=[*reasons, *[f"source:{location}" for location in metadata.source_locations[:8]]],
        text=text,
        normalized_text=normalize_text(text),
        applied=applied,
        candidate_id=_candidate_id(document_name, "footer:default", text),
    )


def _existing_true_footer_present(doc: DocxDocument) -> bool:
    return any(_story_has_visible_content(story) for story in _footer_stories(doc))


def _story_has_visible_content(story: Any) -> bool:
    if any(_compact(paragraph.text) for paragraph in story.paragraphs):
        return True
    return any(_table_text(table) for table in story.tables)


def _footer_stories(doc: DocxDocument) -> list[Any]:
    stories: list[Any] = []
    use_even_footer = bool(getattr(doc.settings, "odd_and_even_pages_header_footer", False))
    for section in doc.sections:
        stories.append(section.footer)
        if section.different_first_page_header_footer:
            stories.append(section.first_page_footer)
        if use_even_footer:
            stories.append(section.even_page_footer)
    return stories


def _header_footer_stories(doc: DocxDocument) -> list[Any]:
    stories: list[Any] = []
    use_even_footer = bool(getattr(doc.settings, "odd_and_even_pages_header_footer", False))
    for section in doc.sections:
        stories.extend((section.header, section.footer))
        if section.different_first_page_header_footer:
            stories.extend((section.first_page_header, section.first_page_footer))
        if use_even_footer:
            stories.extend((section.even_page_header, section.even_page_footer))
    return stories


def _write_real_footer(doc: DocxDocument, metadata: _FooterMetadata, config: dict[str, Any]) -> None:
    sections = list(doc.sections)
    if not sections:
        return

    first_section = sections[0]
    first_section.footer.is_linked_to_previous = False
    _write_footer_story(first_section.footer, metadata, config)

    if bool(config.get("footer_write_first_even_variants", True)):
        if first_section.different_first_page_header_footer:
            first_section.first_page_footer.is_linked_to_previous = False
            _write_footer_story(first_section.first_page_footer, metadata, config)
        if bool(getattr(doc.settings, "odd_and_even_pages_header_footer", False)):
            first_section.even_page_footer.is_linked_to_previous = False
            _write_footer_story(first_section.even_page_footer, metadata, config)

    for section in sections[1:]:
        section.footer.is_linked_to_previous = True
        if section.different_first_page_header_footer:
            section.first_page_footer.is_linked_to_previous = True
        if bool(getattr(doc.settings, "odd_and_even_pages_header_footer", False)):
            section.even_page_footer.is_linked_to_previous = True

    _enable_field_updates(doc)


def _write_footer_story(story: Any, metadata: _FooterMetadata, config: dict[str, Any]) -> None:
    _clear_story_content(story)
    width_inches = float(config.get("footer_table_width_inches", 7.0))
    font_size_pt = float(config.get("footer_font_size_pt", 8))
    table = story.add_table(rows=2, cols=3, width=Inches(width_inches))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    _set_table_borders_none(table)

    for row in table.rows:
        for cell in row.cells:
            cell.width = Inches(width_inches / 3)

    _set_cell_lines(
        table.cell(0, 0),
        [("Versión", True), (metadata.version or "", False)],
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        font_size_pt=font_size_pt,
    )
    _set_cell_lines(
        table.cell(0, 1),
        [(metadata.authorization_date, False)],
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        font_size_pt=font_size_pt,
    )
    _set_page_cell(table.cell(0, 2), metadata, font_size_pt=font_size_pt)
    _set_cell_lines(
        table.cell(1, 0),
        [(metadata.controlled_text, False)],
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        font_size_pt=font_size_pt,
    )
    _set_cell_lines(
        table.cell(1, 1),
        [],
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        font_size_pt=font_size_pt,
    )
    _set_cell_lines(
        table.cell(1, 2),
        [(metadata.next_revision, False)],
        alignment=WD_ALIGN_PARAGRAPH.RIGHT,
        font_size_pt=font_size_pt,
    )


def _clear_story_content(story: Any) -> None:
    for child in list(story._element):
        story._element.remove(child)


def _set_cell_lines(
    cell: Any,
    lines: list[tuple[str, bool]],
    *,
    alignment: WD_ALIGN_PARAGRAPH,
    font_size_pt: float,
) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = alignment
    wrote_line = False
    for text, bold in lines:
        if not text:
            continue
        target_paragraph = paragraph if not wrote_line else cell.add_paragraph()
        target_paragraph.alignment = alignment
        run = target_paragraph.add_run(text)
        run.bold = bold
        run.font.size = Pt(font_size_pt)
        wrote_line = True


def _set_page_cell(cell: Any, metadata: _FooterMetadata, *, font_size_pt: float) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _add_sized_run(paragraph, "Página ", font_size_pt)
    _add_field(paragraph, "PAGE", "1", font_size_pt=font_size_pt)
    _add_sized_run(paragraph, " de ", font_size_pt)
    _add_field(paragraph, "NUMPAGES", metadata.page_total or "1", font_size_pt=font_size_pt)


def _add_sized_run(paragraph: Paragraph, text: str, font_size_pt: float) -> None:
    run = paragraph.add_run(text)
    run.font.size = Pt(font_size_pt)


def _add_field(paragraph: Paragraph, instruction: str, display: str, *, font_size_pt: float) -> None:
    begin_run = paragraph.add_run()
    begin_run.font.size = Pt(font_size_pt)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin_run._r.append(begin)

    instruction_run = paragraph.add_run()
    instruction_run.font.size = Pt(font_size_pt)
    instruction_text = OxmlElement("w:instrText")
    instruction_text.set(qn("xml:space"), "preserve")
    instruction_text.text = f" {instruction} "
    instruction_run._r.append(instruction_text)

    separate_run = paragraph.add_run()
    separate_run.font.size = Pt(font_size_pt)
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run._r.append(separate)

    display_run = paragraph.add_run(display)
    display_run.font.size = Pt(font_size_pt)

    end_run = paragraph.add_run()
    end_run.font.size = Pt(font_size_pt)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)


def _enable_field_updates(doc: DocxDocument) -> None:
    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def _set_table_borders_none(table: Table) -> None:
    table_properties = table._tbl.tblPr
    for existing_borders in table_properties.findall(qn("w:tblBorders")):
        table_properties.remove(existing_borders)
    borders = OxmlElement("w:tblBorders")
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "nil")
        borders.append(border)
    table_properties.append(borders)


def _most_common_text(values: list[str]) -> str:
    compacted = [_compact(value) for value in values if _compact(value)]
    if not compacted:
        return ""
    return Counter(compacted).most_common(1)[0][0]


def _largest_number_text(values: list[str]) -> str:
    numbers = [int(value) for value in values if re.fullmatch(r"\d+", _compact(value))]
    if numbers:
        return str(max(numbers))
    return _most_common_text(values)


def _score_metadata_text(normalized: str, reasons: list[str], config: dict[str, Any]) -> float:
    score = 0.0
    if normalized == "version" or re.fullmatch(r"version\s+\d+(?:\.\d+){0,2}", normalized):
        score = max(score, 0.72)
        reasons.append("footer_label_version")
    if "este es un documento controlado" in normalized:
        score = max(score, 0.88)
        reasons.append("controlled_document_footer")
    if "fecha de autorizacion" in normalized:
        score = max(score, 0.86)
        reasons.append("authorization_date_footer")
    if "proxima revision" in normalized:
        score = max(score, 0.86)
        reasons.append("next_revision_footer")
    if re.fullmatch(r"pagina\s+\d+(?:\s+de\s+\d+)?", normalized):
        score = max(score, 0.88)
        reasons.append("page_number_footer")
    if re.fullmatch(rf"{MONTH_RE}\s+\d{{4}}", normalized):
        score = max(score, 0.5)
        reasons.append("date_only_footer")
    if re.fullmatch(r"[-_=]{4,}", normalized):
        score = max(score, 0.54)
        reasons.append("separator_line")
    if any(re.search(pattern, normalized, re.IGNORECASE) for pattern in config.get("patterns") or []):
        score = max(score, 0.74)
        reasons.append("configured_footer_pattern")
    return score


def _is_metadata_text(normalized: str, config: dict[str, Any]) -> bool:
    if not normalized:
        return False
    if _score_metadata_text(normalized, [], config) >= 0.5:
        return True
    return bool(re.fullmatch(rf"{MONTH_RE}\s+\d{{4}}", normalized))


def _table_label_hits(normalized: str, config: dict[str, Any]) -> int:
    patterns = config.get("table_patterns") or []
    return sum(1 for pattern in patterns if re.search(pattern, normalized, re.IGNORECASE))


def _real_header_footer_texts(doc: DocxDocument) -> set[str]:
    texts: set[str] = set()
    for story in _header_footer_stories(doc):
        for paragraph in story.paragraphs:
            normalized = normalize_text(paragraph.text)
            if normalized:
                texts.add(normalized)
        for table in story.tables:
            texts.update(text for text in _table_cell_norms(table) if text)
    return texts


def _tc_text(tc: Any) -> str:
    parts: list[str] = []
    for child in tc.iterdescendants():
        if child.tag == qn("w:tab"):
            parts.append("\t")
            continue
        if child.tag in {qn("w:noBreakHyphen"), qn("w:softHyphen")}:
            parts.append("-")
            continue
        if child.tag in {qn("w:t"), qn("w:delText"), qn("w:instrText")} and child.text:
            parts.append(child.text)
    return "".join(parts)


def _row_cell_texts(row: Any) -> list[str]:
    try:
        values = [cell.text for cell in row.cells]
    except ValueError:
        values = [_tc_text(tc) for tc in row._tr.tc_lst]
    return [text for text in (_compact(value) for value in values) if text]


def _table_text(table: Table) -> str:
    parts: list[str] = []
    for row in table.rows:
        parts.extend(_row_cell_texts(row))
    return " | ".join(parts)


def _table_signature(table: Table) -> str:
    return "|".join(_table_cell_norms(table))


def _table_cell_norms(table: Table) -> list[str]:
    values: list[str] = []
    for row in table.rows:
        for text in _row_cell_texts(row):
            normalized = normalize_text(text)
            if normalized:
                values.append(normalized)
    return values


def _neighbor_norms(paragraphs: list[_ParagraphInfo], body_index: int, *, radius: int) -> list[str]:
    values: list[str] = []
    for info in paragraphs:
        if info.index == body_index:
            continue
        if abs(info.index - body_index) <= radius:
            values.append(info.normalized)
    return values


def _context_before(paragraphs: list[_ParagraphInfo], body_index: int) -> str:
    previous = [info.text for info in paragraphs if info.index < body_index]
    return previous[-1] if previous else ""


def _context_after(paragraphs: list[_ParagraphInfo], body_index: int) -> str:
    for info in paragraphs:
        if info.index > body_index:
            return info.text
    return ""


def _has_structural_break(paragraph: Paragraph) -> bool:
    paragraph_properties = paragraph._p.pPr
    if paragraph_properties is not None and paragraph_properties.find(qn("w:sectPr")) is not None:
        return True
    for run in paragraph.runs:
        for child in run._r.iterchildren():
            if child.tag in {qn("w:br"), qn("w:lastRenderedPageBreak")}:
                return True
    return False


def _clear_paragraph_text_preserving_breaks(paragraph: Paragraph) -> None:
    for run in paragraph.runs:
        for child in list(run._r.iterchildren()):
            if child.tag in TEXT_CHILD_TAGS:
                run._r.remove(child)


def _remove_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def _remove_table(table: Table) -> None:
    element = table._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


# ---------------------------------------------------------------------------
# Embedded header-table detection and cleanup
# ---------------------------------------------------------------------------

def _is_embedded_header_table(table: Table) -> bool:
    """Return True when the table matches the BHP cover-page header pattern.

    Requires at least 2 of 3 signals:
      1. Image (a:blip) in the first cell of the first row
      2. Approval-row keywords (elaboro/reviso/aprobo) in normalised text
      3. Document-code pattern (p-area-section-NNN)
    """
    normalized = normalize_text(_table_text(table))
    signals = 0

    tbl_elem = table._tbl
    rows = tbl_elem.findall(qn("w:tr"))
    if rows:
        first_row_cells = rows[0].findall(qn("w:tc"))
        if first_row_cells:
            if first_row_cells[0].findall(".//" + _BLIP_QNAME):
                signals += 1

    if _HEADER_APPROVAL_RE.search(normalized):
        signals += 1

    if _HEADER_DOC_CODE_RE.search(normalized):
        signals += 1

    return signals >= 2


def _header_table_infos(table_infos: list[_TableInfo]) -> list[_HeaderTableInfo]:
    result: list[_HeaderTableInfo] = []
    for info in table_infos:
        if _is_embedded_header_table(info.table):
            result.append(
                _HeaderTableInfo(
                    table=info.table,
                    index=info.index,
                    location=f"header_table[{info.index}]",
                    text=info.text,
                    normalized=info.normalized,
                    tbl_element=info.table._tbl,
                )
            )
    return result


def _header_table_candidate(
    info: _HeaderTableInfo,
    *,
    document_name: str,
    occurrence_count: int,
    config: dict[str, Any],
    header_table_position: int,
) -> _ArtifactCandidate | None:
    reasons: list[str] = []
    confidence = 0.0

    if _HEADER_APPROVAL_RE.search(info.normalized):
        confidence = max(confidence, 0.80)
        reasons.append("header_approval_row")

    if _HEADER_DOC_CODE_RE.search(info.normalized):
        confidence = max(confidence, 0.75)
        reasons.append("header_document_code")

    rows = info.tbl_element.findall(qn("w:tr"))
    if rows:
        first_cells = rows[0].findall(qn("w:tc"))
        if first_cells and first_cells[0].findall(".//" + _BLIP_QNAME):
            confidence = max(confidence, 0.72)
            reasons.append("header_logo_image")

    if {"header_approval_row", "header_document_code", "header_logo_image"}.issubset(set(reasons)):
        confidence = max(confidence, 0.95)
        reasons.append("all_header_signals_present")

    if occurrence_count >= 3:
        confidence = min(1.0, confidence + 0.10)
        reasons.append("recurring_header_table")

    confidence = max(0.0, min(1.0, confidence))
    if confidence < float(config.get("min_confidence_review", 0.45)):
        return None

    action = _action_for_header_table(info, confidence, config, reasons, header_table_position=header_table_position)
    record = EmbeddedArtifactRecord(
        document_name=document_name,
        location=info.location,
        block_type="header_table",
        action=action,
        confidence=round(confidence, 3),
        reasons=_unique(reasons),
        text=info.text,
        normalized_text=info.normalized,
        table_index=info.index,
        occurrence_count=occurrence_count,
        recurring_group_id=_hash_id(info.normalized) if occurrence_count > 1 else "",
        candidate_id=_candidate_id(document_name, info.location, info.text),
    )
    return _ArtifactCandidate(record=record, table=info.table, table_index=info.index)


def _action_for_header_table(
    info: _HeaderTableInfo,
    confidence: float,
    config: dict[str, Any],
    reasons: list[str],
    *,
    header_table_position: int,
) -> str:
    action = str(config.get("action", "preview")).lower()
    min_remove = float(config.get("min_confidence_remove", 0.82))
    remove_header_tables = bool(config.get("remove_header_table_artifacts", False))
    front_matter_count = int(config.get("header_table_front_matter_count", 1))

    if header_table_position < front_matter_count:
        reasons.append("protected_front_matter_header_table")
        return "protected"
    if not remove_header_tables:
        reasons.append("header_table_removal_disabled")
        return "review"
    if confidence < min_remove:
        return "review"
    if action == "remove":
        return "remove_table"
    if action == "exclude":
        return "exclude"
    return "would_remove_table"


def _copy_table_element_to_header_part(
    tbl_element: Any,
    source_part: Any,
    header_part: Any,
) -> bool:
    """Deep-copy tbl_element into header_part, rewiring image relationships.

    Returns True on success, False on non-fatal failure (caller skips header
    creation rather than writing a potentially corrupt file).
    """
    from copy import deepcopy

    try:
        tbl_copy = deepcopy(tbl_element)
    except Exception:
        return False

    for blip in tbl_copy.iter(_BLIP_QNAME):
        old_rid = blip.get(_R_EMBED_QNAME)
        if not old_rid:
            continue
        try:
            image_part = source_part.related_parts[old_rid]
            new_rid = header_part.relate_to(image_part, _IMAGE_RELTYPE)
            blip.set(_R_EMBED_QNAME, new_rid)
        except KeyError:
            continue
        except Exception:
            return False

    hdr_elem = header_part.element
    children = list(hdr_elem)
    insert_pos = next(
        (i for i, c in enumerate(children) if c.tag == qn("w:p")),
        len(children),
    )
    hdr_elem.insert(insert_pos, tbl_copy)
    return True


def _write_real_header(
    doc: DocxDocument,
    first_header_tbl_element: Any,
    config: dict[str, Any],
) -> bool:
    sections = list(doc.sections)
    if not sections:
        return False

    first_section = sections[0]
    first_section.header.is_linked_to_previous = False
    header = first_section.header
    _clear_story_content(header)

    success = _copy_table_element_to_header_part(
        first_header_tbl_element,
        source_part=doc.part,
        header_part=header.part,
    )
    if not success:
        return False

    for section in sections[1:]:
        section.header.is_linked_to_previous = True

    _enable_field_updates(doc)
    return True


def _existing_true_header_present(doc: DocxDocument) -> bool:
    return any(_story_has_visible_content(story) for story in _header_stories(doc))


def _header_stories(doc: DocxDocument) -> list[Any]:
    stories: list[Any] = []
    use_even = bool(getattr(doc.settings, "odd_and_even_pages_header_footer", False))
    for section in doc.sections:
        stories.append(section.header)
        if section.different_first_page_header_footer:
            stories.append(section.first_page_header)
        if use_even:
            stories.append(section.even_page_header)
    return stories


def _make_header_record(
    *,
    document_name: str,
    action: str,
    source_location: str,
    reasons: list[str],
    applied: bool,
    confidence: float = 1.0,
) -> EmbeddedArtifactRecord:
    return EmbeddedArtifactRecord(
        document_name=document_name,
        location="header:default",
        block_type="header",
        action=action,
        confidence=confidence,
        reasons=[*reasons, f"source:{source_location}"],
        text="",
        normalized_text="",
        applied=applied,
        candidate_id=_candidate_id(document_name, "header:default", action),
    )


def _header_reconstruction_record(
    doc: DocxDocument,
    candidates: list[_ArtifactCandidate],
    *,
    document_name: str,
    config: dict[str, Any],
    apply_changes: bool,
) -> tuple[EmbeddedArtifactRecord | None, bool]:
    if not bool(config.get("write_real_header", True)):
        return None, False
    if not bool(config.get("remove_header_table_artifacts", False)):
        return None, False

    source_candidate = next(
        (
            c
            for c in candidates
            if c.record.block_type == "header_table"
            and c.record.action == "protected"
            and c.table is not None
        ),
        None,
    )
    if source_candidate is None:
        return None, False

    removable = [
        c
        for c in candidates
        if c.record.block_type == "header_table"
        and c.record.action in {"remove_table", "would_remove_table"}
    ]
    if not removable:
        return None, False

    if _existing_true_header_present(doc) and not bool(config.get("overwrite_existing_header", False)):
        return (
            _make_header_record(
                document_name=document_name,
                action="header_protected_existing",
                source_location=source_candidate.record.location,
                reasons=["existing_true_header_detected", "overwrite_existing_header_disabled"],
                applied=False,
            ),
            False,
        )

    action = "write_header" if apply_changes else "would_write_header"
    applied = False
    if apply_changes:
        applied = _write_real_header(doc, source_candidate.table._tbl, config)
        if not applied:
            action = "header_write_failed"
        else:
            source_candidate.record.action = "move_to_header"
            source_candidate.record.reasons = _unique(
                [*source_candidate.record.reasons, "moved_to_real_header"]
            )
    else:
        source_candidate.record.action = "would_move_to_header"
        source_candidate.record.reasons = _unique(
            [*source_candidate.record.reasons, "would_move_to_real_header"]
        )
    return (
        _make_header_record(
            document_name=document_name,
            action=action,
            source_location=source_candidate.record.location,
            reasons=["reconstructed_real_header", "table_with_image_copied"],
            applied=applied,
        ),
        applied,
    )


def _alignment_name(paragraph: Paragraph) -> str:
    if paragraph.alignment is None:
        return ""
    return getattr(paragraph.alignment, "name", str(paragraph.alignment))


def _compact(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


def _unique(values: list[str]) -> list[str]:
    return list(dict.fromkeys(values))


def _hash_id(value: str) -> str:
    return hashlib.sha1(value.encode("utf-8")).hexdigest()[:12]


def _candidate_id(document_name: str, location: str, text: str) -> str:
    return _hash_id(f"{document_name}\0{location}\0{text}")