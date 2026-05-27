from __future__ import annotations

import json
from pathlib import Path

from docx import Document


PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUTPUT_PATH = PROJECT_ROOT / "data" / "input" / "mock_reglas.docx"
SAMPLES_PATH = PROJECT_ROOT / "reports" / "referencia_operador_samples.jsonl"


def _set_cell_text(cell, text: str) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.text = text


def _add_row(table, values: list[str]) -> None:
    cells = table.add_row().cells
    for index, value in enumerate(values):
        _set_cell_text(cells[index], value)


def _load_reference_examples(limit: int = 3) -> list[str]:
    if not SAMPLES_PATH.exists():
        return []
    examples: list[str] = []
    with SAMPLES_PATH.open("r", encoding="utf-8") as handle:
        for line in handle:
            if len(examples) >= limit:
                break
            try:
                record = json.loads(line)
            except json.JSONDecodeError:
                continue
            text = str(record.get("text", "")).strip()
            lowered = text.lower()
            if text and "supervisor" not in lowered and "one-up" not in lowered:
                examples.append(text[:260])
    return examples


def main() -> int:
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading("Mock procedimiento reglas", level=1)
    doc.add_paragraph(
        "Documento editable de prueba para validar cambios de cargos responsables en tablas y parrafos."
    )

    doc.add_heading("1. Control operacional", level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["Actividad", "Responsable", "Instruccion", "Resultado esperado"]
    for index, header in enumerate(headers):
        _set_cell_text(table.rows[0].cells[index], header)

    _add_row(
        table,
        [
            "Inspeccion inicial",
            "operador",
            "El operador debe verificar el estado de las correas antes del arranque.",
            "Registro completado.",
        ],
    )
    _add_row(
        table,
        [
            "Cambio de turno",
            "operadora de planta",
            "La operadora de planta informa desviaciones al jefe de turno.",
            "Desviaciones comunicadas.",
        ],
    )
    _add_row(
        table,
        [
            "Ronda de terreno",
            "operadores de terreno",
            "Los operadores de terreno deben aislar el sector y reportar la condicion.",
            "Sector controlado.",
        ],
    )
    _add_row(
        table,
        [
            "Aviso operacional",
            "supervisor",
            "Ante alarma critica se debe informar al operador de sala para coordinar acciones.",
            "Aviso registrado.",
        ],
    )
    _add_row(
        table,
        [
            "Coordinacion segura",
            "supervisor de operaciones",
            "El supervisor de operaciones debe coordinar el ingreso al area restringida.",
            "Ingreso controlado.",
        ],
    )
    _add_row(
        table,
        [
            "Comunicacion de evento",
            "operador planta",
            "El operador planta debe informar al supervisor de turno cualquier desviacion critica.",
            "Evento comunicado.",
        ],
    )
    _add_row(
        table,
        [
            "Caso supervisor ya expandido",
            "supervisor o Ejecutivos del Área",
            "El supervisor o Ejecutivos del Área autoriza la continuidad de la tarea.",
            "Sin duplicados.",
        ],
    )
    _add_row(
        table,
        [
            "Supervisor exento",
            "Supervisor Sala de Control",
            "El Supervisor Sala de Control monitorea las variables desde CAS.",
            "Sin cambio por excepcion.",
        ],
    )
    _add_row(
        table,
        [
            "Caso ya expandido",
            "operador o personal designado por minera Spence",
            "El operador o personal designado por minera Spence confirma la disponibilidad del equipo.",
            "Sin duplicados.",
        ],
    )
    _add_row(
        table,
        [
            "Alternativa existente",
            "operador o mantenedor",
            "El operador o mantenedor puede detener el equipo si observa condicion insegura.",
            "Se envia a revision, sin duplicar.",
        ],
    )
    _add_row(
        table,
        [
            "Texto sin accion",
            "Perfil del operador",
            "Perfil del operador: competencias generales del cargo.",
            "Sin cambio automatico.",
        ],
    )

    split_row = table.add_row().cells
    _set_cell_text(split_row[0], "Runs divididos")
    p = split_row[1].paragraphs[0]
    p.text = ""
    p.add_run("El ")
    p.add_run("Oper").bold = True
    p.add_run("ador")
    p.add_run(" de planta valida el bloqueo local.")
    _set_cell_text(split_row[2], "Caso para preservar formato de runs.")
    _set_cell_text(split_row[3], "Formato conservado.")

    doc.add_heading("2. Apariciones fuera de tabla", level=2)
    doc.add_paragraph("La operadora debe cerrar la actividad en el sistema de turno.")
    doc.add_paragraph("Los operadores de planta revisan la comunicacion radial antes de iniciar.")
    doc.add_paragraph("La supervisora de turno revisa las condiciones generales antes de liberar el equipo.")
    doc.add_paragraph("El jefe de área debe validar que el permiso de trabajo este vigente.")

    reference_examples = _load_reference_examples()
    if reference_examples:
        doc.add_heading("3. Extractos de referencia", level=2)
        ref_table = doc.add_table(rows=1, cols=2)
        ref_table.style = "Table Grid"
        _set_cell_text(ref_table.rows[0].cells[0], "Origen")
        _set_cell_text(ref_table.rows[0].cells[1], "Texto adaptado")
        for index, example in enumerate(reference_examples, start=1):
            _add_row(ref_table, [f"referencia.xlsx #{index}", example])

    doc.save(str(OUTPUT_PATH))
    print(f"Wrote {OUTPUT_PATH.relative_to(PROJECT_ROOT)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())